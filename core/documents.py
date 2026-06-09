"""
Формирование документов из приложения в форматах .docx и .xlsx.

Используются библиотеки python-docx (Word) и openpyxl (Excel). Документы
собираются «на лету» по данным БД и отдаются пользователю как файл на скачивание.
"""
import io

from django.http import HttpResponse
from django.utils import timezone

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill, Border, Side

from . import analytics
from .models import Application, Candidate, Vacancy

AUTHOR_FOOTER = "Сформировано в ИС «UnitHire». Автор системы: Серебренников Д. В."


# ---------------------------------------------------------------------------
#  Word: карточка кандидата
# ---------------------------------------------------------------------------
def candidate_card_docx(candidate: Candidate) -> HttpResponse:
    """Сформировать карточку кандидата в формате .docx."""
    doc = Document()
    # Базовый стиль документа — Times New Roman 12.
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    title = doc.add_heading("Карточка кандидата", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(candidate.full_name())
    run.bold = True
    run.font.size = Pt(14)

    # Таблица основных реквизитов.
    rows = [
        ("Email", candidate.email),
        ("Телефон", candidate.phone or "—"),
        ("Город", candidate.city or "—"),
        ("Грейд", candidate.get_grade_display()),
        ("Опыт, лет", str(candidate.experience_years)),
        ("Желаемая зарплата, ₽", f"{candidate.desired_salary:,}".replace(",", " ")),
        ("Источник", candidate.source.name),
        ("Дата добавления", candidate.created_at.strftime("%d.%m.%Y")),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = str(value)

    # Навыки.
    doc.add_heading("Навыки", level=1)
    skills = candidate.candidateskill_set.select_related("skill")
    if skills:
        for cs in skills:
            doc.add_paragraph(
                f"{cs.skill.name} — {cs.get_level_display()}", style="List Bullet")
    else:
        doc.add_paragraph("Навыки не указаны.")

    # Отклики кандидата.
    doc.add_heading("Отклики на вакансии", level=1)
    apps = candidate.applications.select_related("vacancy", "stage")
    if apps:
        t = doc.add_table(rows=1, cols=3)
        t.style = "Light List Accent 1"
        hdr = t.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Вакансия", "Этап", "Статус"
        for a in apps:
            c = t.add_row().cells
            c[0].text = a.vacancy.title
            c[1].text = a.stage.name
            c[2].text = a.get_status_display()
    else:
        doc.add_paragraph("Откликов нет.")

    if candidate.summary:
        doc.add_heading("О кандидате", level=1)
        doc.add_paragraph(candidate.summary)

    footer = doc.add_paragraph()
    footer.add_run(
        f"\n{AUTHOR_FOOTER}\nДата формирования: "
        f"{timezone.localtime().strftime('%d.%m.%Y %H:%M')}"
    ).italic = True

    return _docx_response(doc, f"candidate_{candidate.pk}.docx")


def _docx_response(doc: Document, filename: str) -> HttpResponse:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    resp = HttpResponse(
        buf.read(),
        content_type="application/vnd.openxmlformats-officedocument."
                     "wordprocessingml.document")
    resp["Content-Disposition"] = f'attachment; filename="{filename}"'
    return resp


# ---------------------------------------------------------------------------
#  Excel: аналитические отчёты
# ---------------------------------------------------------------------------
_HEADER_FILL = PatternFill("solid", fgColor="2563EB")
_HEADER_FONT = Font(color="FFFFFF", bold=True)
_THIN = Side(style="thin", color="BBBBBB")
_BORDER = Border(left=_THIN, right=_THIN, top=_THIN, bottom=_THIN)


def _write_sheet(ws, headers, rows):
    """Заполнить лист Excel: шапка + строки данных + оформление."""
    ws.append(headers)
    for col, _ in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col)
        cell.fill = _HEADER_FILL
        cell.font = _HEADER_FONT
        cell.alignment = Alignment(horizontal="center", vertical="center",
                                   wrap_text=True)
        cell.border = _BORDER
    for r in rows:
        ws.append(r)
    # Границы и автоширина по содержимому.
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, max_col=len(headers)):
        for cell in row:
            cell.border = _BORDER
    for col_idx, header in enumerate(headers, start=1):
        width = max(len(str(header)), 12)
        for row in ws.iter_rows(min_row=2, min_col=col_idx, max_col=col_idx):
            for cell in row:
                width = max(width, len(str(cell.value or "")))
        ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = \
            min(width + 2, 45)


def analytics_xlsx() -> HttpResponse:
    """Выгрузить сводный аналитический отчёт (несколько листов) в Excel."""
    wb = Workbook()

    # Лист 1 — воронка подбора.
    ws = wb.active
    ws.title = "Воронка"
    _write_sheet(ws, ["Этап воронки", "Заявок", "Конверсия, %"],
                 [[r["name"], r["applications"], r["conversion"]]
                  for r in analytics.funnel_report()])

    # Лист 2 — эффективность источников.
    ws2 = wb.create_sheet("Источники")
    _write_sheet(ws2,
                 ["Источник", "Кандидатов", "Откликов", "Нанято",
                  "Конверсия, %", "Стоимость найма, ₽"],
                 [[r["source"], r["candidates"], r["applications"], r["hired"],
                   r["conversion"], r["cost_per_hire"] if r["cost_per_hire"] else "—"]
                  for r in analytics.source_efficiency_report()])

    # Лист 3 — время закрытия вакансий.
    ws3 = wb.create_sheet("Time-to-hire")
    _write_sheet(ws3,
                 ["Отдел", "Закрыто вакансий", "Среднее, дней",
                  "Минимум", "Максимум"],
                 [[r["department"], r["closed"], r["avg_days"],
                   r["min_days"], r["max_days"]]
                  for r in analytics.time_to_hire_report()])

    # Лист 4 — загрузка рекрутёров.
    ws4 = wb.create_sheet("Загрузка рекрутёров")
    _write_sheet(ws4,
                 ["Рекрутёр", "Открытых вакансий", "Активных заявок",
                  "Собеседований", "Нанято"],
                 [[r["recruiter"], r["open_vacancies"], r["active_applications"],
                   r["interviews"], r["hired"]]
                  for r in analytics.recruiter_load_report()])

    return _xlsx_response(wb, "hr_analytics.xlsx")


def vacancies_xlsx(queryset=None) -> HttpResponse:
    """Выгрузить список вакансий в Excel."""
    queryset = queryset if queryset is not None else Vacancy.objects.all()
    wb = Workbook()
    ws = wb.active
    ws.title = "Вакансии"
    rows = []
    for v in queryset.select_related("department", "recruiter"):
        rows.append([
            v.title, v.department.name, v.get_grade_display(),
            v.get_status_display(), v.salary_min, v.salary_max,
            v.recruiter.short_name() if v.recruiter else "—",
            v.applications.count(),
            v.opened_at.strftime("%d.%m.%Y"),
            v.closed_at.strftime("%d.%m.%Y") if v.closed_at else "—",
        ])
    _write_sheet(ws,
                 ["Вакансия", "Отдел", "Грейд", "Статус", "Зарплата от",
                  "Зарплата до", "Рекрутёр", "Откликов", "Открыта", "Закрыта"],
                 rows)
    return _xlsx_response(wb, "vacancies.xlsx")


def _xlsx_response(wb: Workbook, filename: str) -> HttpResponse:
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    resp = HttpResponse(
        buf.read(),
        content_type="application/vnd.openxmlformats-officedocument."
                     "spreadsheetml.sheet")
    resp["Content-Disposition"] = f'attachment; filename="{filename}"'
    return resp
