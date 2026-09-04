"""Извлечение текста и стажа работы из файла резюме.

Поддерживаются форматы .pdf, .docx и .txt. Если формат не поддержан
или файл повреждён, возвращается пустая строка: подсистема отбора
должна сообщить рекрутёру о невозможности разбора, а не прервать
обработку отклика.
"""
from __future__ import annotations

import logging
import re
from pathlib import Path

logger = logging.getLogger(__name__)

# Стаж в резюме пишут по-разному: «опыт работы 5 лет», «стаж: 3 года»,
# «7+ лет опыта». Разбираются все три написания.
EXPERIENCE_PATTERNS = [
    re.compile(r"опыт\D{0,20}?(\d{1,2})\s*(?:\+\s*)?(?:год|года|лет)",
               re.IGNORECASE),
    re.compile(r"стаж\D{0,20}?(\d{1,2})\s*(?:\+\s*)?(?:год|года|лет)",
               re.IGNORECASE),
    re.compile(r"(\d{1,2})\s*\+?\s*(?:год|года|лет)\D{0,15}?опыт",
               re.IGNORECASE),
]


def extract_text(file_path: str) -> str:
    """Возвращает текст файла резюме."""
    path = Path(file_path)
    if not path.exists():
        logger.warning("Файл резюме не найден: %s", file_path)
        return ""

    suffix = path.suffix.lower()
    try:
        if suffix == ".docx":
            return _extract_docx(path)
        if suffix == ".pdf":
            return _extract_pdf(path)
        if suffix == ".txt":
            return path.read_text(encoding="utf-8", errors="ignore")
    except Exception as exc:
        logger.error("Ошибка разбора файла %s: %s", file_path, exc)
        return ""

    logger.info("Формат %s не поддержан разборщиком", suffix)
    return ""


def _extract_docx(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    # Часть кандидатов оформляет резюме таблицей, поэтому ячейки тоже
    # попадают в текст.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts)


def _extract_pdf(path: Path) -> str:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        logger.info("Разбор PDF недоступен: библиотека не установлена")
        return ""

    with fitz.open(str(path)) as document:
        return "\n".join(page.get_text() for page in document)


def extract_experience_years(text: str) -> float | None:
    """Определяет стаж работы по тексту резюме.

    Возвращает наибольшее найденное значение: кандидаты нередко
    указывают и общий стаж, и стаж по отдельным местам работы.
    """
    if not text:
        return None
    values = []
    for pattern in EXPERIENCE_PATTERNS:
        for match in pattern.finditer(text):
            try:
                value = float(match.group(1))
            except (TypeError, ValueError):
                continue
            # Значения свыше 50 лет - почти наверняка не про стаж.
            if 0 < value <= 50:
                values.append(value)
    return max(values) if values else None
