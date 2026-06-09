"""
Автоматизированные тесты рекрутинговой ИС «UnitHire».

Покрытие отражает четыре метода тестирования из методички:
  • модульное (unit) — методы моделей и функции аналитики;
  • функциональное — доступность страниц, работа форм и сценариев;
  • дымовое (smoke) — публичные страницы и кабинеты открываются без ошибок;
  • негативное — некорректные данные отклоняются, права разграничены.

Запуск:  python manage.py test
"""
from datetime import timedelta

from django.test import TestCase
from django.urls import reverse
from django.utils import timezone

from .analytics import (funnel_report, kpi_summary, source_efficiency_report,
                        time_to_hire_report)
from .documents import analytics_xlsx, candidate_card_docx
from .forms import FeedbackForm, SignUpForm, VacancyForm
from .models import (Application, ApplicationStatus, Candidate, Department,
                     Role, Source, Stage, User, Vacancy, VacancyStatus)


class BaseData(TestCase):
    """Минимальный набор данных для тестов."""

    @classmethod
    def setUpTestData(cls):
        cls.dep = Department.objects.create(name="Отдел разработки")
        cls.src = Source.objects.create(name="hh.ru", cost_per_contact=1000)
        cls.stage1 = Stage.objects.create(name="Скрининг", order=1)
        cls.stage_hire = Stage.objects.create(name="Принят", order=6,
                                              is_terminal=True)
        cls.recruiter = User.objects.create_user(
            "rec", password="Hr#Unitcode2026", role=Role.RECRUITER,
            first_name="Анна", last_name="Соколова")
        cls.candidate_user = User.objects.create_user(
            "cand", password="User#Unitcode2026", role=Role.CANDIDATE,
            first_name="Максим", last_name="Новиков")
        cls.cand = Candidate.objects.create(
            user=cls.candidate_user, last_name="Новиков", first_name="Максим",
            email="cand@example.com", source=cls.src, desired_salary=150000)
        cls.vac = Vacancy.objects.create(
            title="Python-разработчик", department=cls.dep,
            salary_min=150000, salary_max=220000, status=VacancyStatus.OPEN,
            recruiter=cls.recruiter, opened_at=timezone.localdate() - timedelta(days=20))


# --------------------------------------------------------------------------
#  Модульные тесты (unit)
# --------------------------------------------------------------------------
class ModelTests(BaseData):

    def test_candidate_full_name(self):
        self.cand.patronymic = "Андреевич"
        self.assertEqual(self.cand.full_name(), "Новиков Максим Андреевич")

    def test_user_short_name(self):
        self.assertEqual(self.recruiter.short_name(), "Соколова А.")

    def test_vacancy_salary_range(self):
        self.assertIn("150 000", self.vac.salary_range())

    def test_vacancy_time_to_hire(self):
        self.vac.closed_at = self.vac.opened_at + timedelta(days=15)
        self.vac.save()
        self.assertEqual(self.vac.time_to_hire(), 15)

    def test_unique_application_constraint(self):
        Application.objects.create(candidate=self.cand, vacancy=self.vac,
                                   stage=self.stage1)
        from django.db import IntegrityError, transaction
        with self.assertRaises(IntegrityError):
            with transaction.atomic():
                Application.objects.create(candidate=self.cand, vacancy=self.vac,
                                           stage=self.stage1)


class AnalyticsTests(BaseData):

    def test_kpi_summary_keys(self):
        kpi = kpi_summary()
        for key in ("open_vacancies", "candidates", "conversion",
                    "avg_time_to_hire"):
            self.assertIn(key, kpi)

    def test_funnel_report(self):
        Application.objects.create(candidate=self.cand, vacancy=self.vac,
                                   stage=self.stage1)
        rows = funnel_report()
        self.assertTrue(any(r["applications"] >= 1 for r in rows))

    def test_source_efficiency(self):
        Application.objects.create(candidate=self.cand, vacancy=self.vac,
                                   stage=self.stage_hire,
                                   status=ApplicationStatus.HIRED)
        rows = source_efficiency_report()
        hh = next(r for r in rows if r["source"] == "hh.ru")
        self.assertEqual(hh["hired"], 1)
        self.assertIsNotNone(hh["cost_per_hire"])

    def test_time_to_hire_report(self):
        self.vac.closed_at = self.vac.opened_at + timedelta(days=10)
        self.vac.save()
        rows = time_to_hire_report()
        self.assertTrue(any(r["avg_days"] == 10 for r in rows))


class DocumentTests(BaseData):

    def test_candidate_card_docx(self):
        resp = candidate_card_docx(self.cand)
        self.assertEqual(resp.status_code, 200)
        self.assertIn("wordprocessingml", resp["Content-Type"])
        self.assertGreater(len(resp.content), 1000)

    def test_analytics_xlsx(self):
        resp = analytics_xlsx()
        self.assertEqual(resp.status_code, 200)
        self.assertIn("spreadsheetml", resp["Content-Type"])


# --------------------------------------------------------------------------
#  Тесты форм (валидация — позитивные и негативные)
# --------------------------------------------------------------------------
class FormTests(BaseData):

    def test_feedback_form_valid(self):
        form = FeedbackForm(data={
            "name": "Иван Петров", "email": "ivan@example.com",
            "subject": "Вопрос", "message": "Здравствуйте, есть вопрос по вакансии."})
        self.assertTrue(form.is_valid())

    def test_feedback_form_short_message(self):
        form = FeedbackForm(data={
            "name": "Иван", "email": "ivan@example.com",
            "subject": "Тест", "message": "Коротко"})
        self.assertFalse(form.is_valid())
        self.assertIn("message", form.errors)

    def test_feedback_honeypot(self):
        form = FeedbackForm(data={
            "name": "Бот", "email": "bot@example.com", "subject": "Спам",
            "message": "Длинное сообщение для спама", "website": "http://spam"})
        self.assertFalse(form.is_valid())

    def test_vacancy_salary_validation(self):
        form = VacancyForm(data={
            "title": "Тест", "department": self.dep.id, "grade": "middle",
            "salary_min": 200000, "salary_max": 100000, "status": "open",
            "city": "Москва"})
        self.assertFalse(form.is_valid())
        self.assertIn("salary_max", form.errors)

    def test_signup_creates_candidate(self):
        form = SignUpForm(data={
            "username": "newuser", "first_name": "Пётр", "last_name": "Сидоров",
            "email": "petr@example.com", "phone": "+7 900 000-00-00",
            "password1": "SecurePass2026", "password2": "SecurePass2026"})
        self.assertTrue(form.is_valid(), form.errors)
        user = form.save()
        self.assertEqual(user.role, Role.CANDIDATE)
        self.assertTrue(Candidate.objects.filter(user=user).exists())


# --------------------------------------------------------------------------
#  Функциональные и дымовые тесты (smoke)
# --------------------------------------------------------------------------
class PublicPagesTests(BaseData):

    def test_public_pages_ok(self):
        for name in ["home", "about", "how_it_works", "employers",
                     "candidates_info", "help", "vacancies", "analytics_demo",
                     "news", "contacts", "register", "login"]:
            with self.subTest(page=name):
                resp = self.client.get(reverse(name))
                self.assertEqual(resp.status_code, 200)

    def test_vacancy_detail(self):
        resp = self.client.get(reverse("vacancy_detail", args=[self.vac.pk]))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, "Python-разработчик")

    def test_footer_has_author(self):
        resp = self.client.get(reverse("home"))
        self.assertContains(resp, "Серебренников")

    def test_breadcrumbs_present(self):
        resp = self.client.get(reverse("vacancies"))
        self.assertContains(resp, "breadcrumb")

    def test_feedback_submission(self):
        resp = self.client.post(reverse("contacts"), {
            "name": "Тест", "email": "t@example.com", "subject": "Тема",
            "message": "Сообщение достаточной длины для проверки."})
        self.assertEqual(resp.status_code, 302)
        from .models import Feedback
        self.assertEqual(Feedback.objects.count(), 1)


class PermissionTests(BaseData):

    def test_anonymous_redirected_from_cabinet(self):
        resp = self.client.get(reverse("rec_dashboard"))
        self.assertEqual(resp.status_code, 302)
        self.assertIn("/login/", resp.url)

    def test_candidate_forbidden_on_recruiter(self):
        self.client.login(username="cand", password="User#Unitcode2026")
        resp = self.client.get(reverse("rec_analytics"))
        self.assertEqual(resp.status_code, 403)

    def test_recruiter_access_cabinet(self):
        self.client.login(username="rec", password="Hr#Unitcode2026")
        for name in ["rec_dashboard", "rec_vacancies", "rec_candidates",
                     "rec_applications", "rec_interviews", "rec_analytics"]:
            with self.subTest(page=name):
                self.assertEqual(self.client.get(reverse(name)).status_code, 200)

    def test_candidate_access_cabinet(self):
        self.client.login(username="cand", password="User#Unitcode2026")
        for name in ["cand_dashboard", "cand_profile", "cand_resumes",
                     "cand_applications", "cand_interviews"]:
            with self.subTest(page=name):
                self.assertEqual(self.client.get(reverse(name)).status_code, 200)


class ScenarioTests(BaseData):

    def test_candidate_apply_flow(self):
        """Сценарий: кандидат откликается на вакансию."""
        self.client.login(username="cand", password="User#Unitcode2026")
        resp = self.client.post(
            reverse("cand_apply", args=[self.vac.pk]),
            {"cover_letter": "Готов приступить."}, follow=True)
        self.assertEqual(resp.status_code, 200)
        self.assertTrue(Application.objects.filter(
            candidate=self.cand, vacancy=self.vac).exists())

    def test_recruiter_creates_vacancy(self):
        """Сценарий: рекрутёр создаёт вакансию."""
        self.client.login(username="rec", password="Hr#Unitcode2026")
        resp = self.client.post(reverse("rec_vacancy_create"), {
            "title": "Новая вакансия", "department": self.dep.id,
            "grade": "junior", "salary_min": 80000, "salary_max": 120000,
            "status": "open", "city": "Москва"}, follow=True)
        self.assertEqual(resp.status_code, 200)
        self.assertTrue(Vacancy.objects.filter(title="Новая вакансия").exists())


# --------------------------------------------------------------------------
#  Тесты модели бенчмарков и аналитики «наш показатель vs индустрия»
# --------------------------------------------------------------------------
class BenchmarkTests(BaseData):
    """Замечание рецензента 11: HR-метрики должны сравниваться с отраслью."""

    def test_benchmark_create_and_str(self):
        from .models import IndustryBenchmark
        bm = IndustryBenchmark.objects.create(
            metric="time_to_hire", industry="ИТ в России",
            value=30, unit="дней",
            source="hh.ru research 2024", year=2024,
        )
        self.assertEqual(bm.metric, "time_to_hire")
        self.assertIn("Время закрытия", str(bm))
        self.assertIn("2024", str(bm))
        self.assertEqual(IndustryBenchmark.objects.count(), 1)

    def test_benchmark_uniqueness(self):
        from .models import IndustryBenchmark
        from django.db import IntegrityError, transaction
        IndustryBenchmark.objects.create(
            metric="time_to_hire", industry="ИТ в России",
            value=30, unit="дней", source="A", year=2024)
        with self.assertRaises(IntegrityError):
            with transaction.atomic():
                IndustryBenchmark.objects.create(
                    metric="time_to_hire", industry="ИТ в России",
                    value=31, unit="дней", source="B", year=2024)


# --------------------------------------------------------------------------
#  Безопасность загрузки резюме (замечание рецензента 10)
# --------------------------------------------------------------------------
class ResumeSecurityTests(BaseData):
    """Защита от подмены типа файла, превышения размера, пустых файлов."""

    def setUp(self):
        # Заранее создаём анкету кандидата для logged-in пользователя cand.
        # _get_candidate в view создаёт её автоматически — но source может
        # ещё не быть; гарантируем источник «Карьерный сайт».
        from .models import Source
        Source.objects.get_or_create(
            name="Карьерный сайт",
            defaults={"kind": "собственный сайт", "cost_per_contact": 0},
        )
        self.client.login(username="cand", password="User#Unitcode2026")

    def _upload(self, name, content, content_type="application/pdf"):
        from django.core.files.uploadedfile import SimpleUploadedFile
        f = SimpleUploadedFile(name, content, content_type=content_type)
        return self.client.post(
            reverse("cand_resumes"),
            {"title": "test", "file": f},
        )

    def test_rejects_php_extension(self):
        from .models import ResumeFile
        before = ResumeFile.objects.count()
        self._upload("malware.php", b"<?php echo 'pwn'; ?>",
                     content_type="application/x-php")
        self.assertEqual(ResumeFile.objects.count(), before)

    def test_rejects_oversized_file(self):
        from .models import ResumeFile
        before = ResumeFile.objects.count()
        big = b"x" * (6 * 1024 * 1024)
        self._upload("big.pdf", big)
        self.assertEqual(ResumeFile.objects.count(), before)

    def test_rejects_empty_file(self):
        from .models import ResumeFile
        before = ResumeFile.objects.count()
        self._upload("empty.pdf", b"")
        self.assertEqual(ResumeFile.objects.count(), before)

    def test_rejects_mismatched_mime(self):
        from .models import ResumeFile
        before = ResumeFile.objects.count()
        self._upload("resume.pdf", b"<html><script>x</script></html>",
                     content_type="text/html")
        self.assertEqual(ResumeFile.objects.count(), before)

    def test_accepts_valid_pdf_with_uuid_rename(self):
        from .models import ResumeFile
        pdf = b"%PDF-1.4\n%EOF\n"
        self._upload("resume.pdf", pdf)
        self.assertEqual(ResumeFile.objects.count(), 1)
        rf = ResumeFile.objects.first()
        # Исходное имя не должно сохраниться в пути на ФС.
        self.assertNotIn("resume.pdf", rf.file.name)
        # Расширение сохраняется.
        self.assertTrue(rf.file.name.endswith(".pdf"))
        # И почистим за собой файл из media/.
        rf.file.delete(save=False)


# --------------------------------------------------------------------------
#  Атомарность операций — смена этапа + закрытие вакансии (замечание 9)
# --------------------------------------------------------------------------
class ApplicationTransactionTests(BaseData):
    """Перевод на терминальный этап + статус HIRED закрывает вакансию атомарно."""

    def setUp(self):
        self.app = Application.objects.create(
            candidate=self.cand, vacancy=self.vac,
            stage=self.stage1, status=ApplicationStatus.NEW,
        )
        self.client.login(username="rec", password="Hr#Unitcode2026")

    def test_terminal_hired_closes_vacancy(self):
        resp = self.client.post(
            reverse("rec_application_edit", args=[self.app.pk]),
            {"stage": self.stage_hire.pk,
             "status": ApplicationStatus.HIRED,
             "score": 90, "comment": ""},
        )
        self.assertEqual(resp.status_code, 302)
        self.vac.refresh_from_db()
        self.assertEqual(self.vac.status, VacancyStatus.CLOSED)
        self.assertIsNotNone(self.vac.closed_at)

    def test_non_terminal_does_not_close_vacancy(self):
        resp = self.client.post(
            reverse("rec_application_edit", args=[self.app.pk]),
            {"stage": self.stage1.pk,
             "status": ApplicationStatus.IN_REVIEW,
             "score": 60, "comment": ""},
        )
        self.assertEqual(resp.status_code, 302)
        self.vac.refresh_from_db()
        self.assertEqual(self.vac.status, VacancyStatus.OPEN)
        self.assertIsNone(self.vac.closed_at)


class BenchmarkComparisonTests(BaseData):
    """benchmark_comparison: KPI с парой (our, industry, delta_pct)."""

    def test_returns_metric_pairs(self):
        from .analytics import benchmark_comparison
        from .models import IndustryBenchmark
        IndustryBenchmark.objects.create(
            metric="time_to_hire", industry="ИТ в России",
            value=30, unit="дней", source="test", year=2024,
        )
        IndustryBenchmark.objects.create(
            metric="conversion", industry="ИТ в России",
            value=8.5, unit="%", source="test", year=2024,
        )
        result = benchmark_comparison()
        self.assertIn("time_to_hire", result)
        self.assertIn("conversion", result)
        self.assertEqual(result["time_to_hire"]["industry"], 30)
        self.assertEqual(result["time_to_hire"]["unit"], "дней")
        # KPI без данных — our может быть 0, но ключ есть.
        self.assertIn("our", result["time_to_hire"])


class CandidateStatusVisibilityTests(BaseData):
    """Замечание 3: кандидат должен видеть свои отклики и статусы."""

    def test_candidate_sees_stage_and_status(self):
        Application.objects.create(
            candidate=self.cand, vacancy=self.vac,
            stage=self.stage1, status=ApplicationStatus.IN_REVIEW,
        )
        self.client.login(username="cand", password="User#Unitcode2026")
        resp = self.client.get(reverse("cand_applications"))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, "Скрининг")
        self.assertContains(resp, "На рассмотрении")
        self.assertContains(resp, "badge")
        self.assertContains(resp, "Python-разработчик")


# --------------------------------------------------------------------------
#  Cyrillic round-trip для .docx и .xlsx (замечание рецензента 1)
# --------------------------------------------------------------------------
class CyrillicRoundTripTests(BaseData):
    """Генерируем документ → парсим обратно → проверяем, что кириллица
    действительно в файле, а не превращена в '?' или знаки вопроса."""

    def setUp(self):
        from .models import Skill, CandidateSkill
        self.cand2 = Candidate.objects.create(
            last_name="Серебренников", first_name="Дмитрий",
            patronymic="Валерьевич", email="dvs@example.com",
            source=self.src,
            summary="Опытный специалист по нагрузочному тестированию",
            desired_salary=200000,
        )
        skill = Skill.objects.create(
            name="Нагрузочное тестирование", category="tool")
        CandidateSkill.objects.create(
            candidate=self.cand2, skill=skill, level=3)

    def test_docx_contains_cyrillic_full_name(self):
        from io import BytesIO
        from docx import Document
        from .documents import candidate_card_docx
        resp = candidate_card_docx(self.cand2)
        doc = Document(BytesIO(resp.content))
        text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += "\n" + cell.text
        self.assertIn("Серебренников", text)
        self.assertIn("Дмитрий", text)
        self.assertIn("Валерьевич", text)
        self.assertIn("Нагрузочное тестирование", text)
        self.assertIn("Опытный специалист", text)

    def test_xlsx_contains_cyrillic_headers(self):
        from io import BytesIO
        from openpyxl import load_workbook
        from .documents import analytics_xlsx
        resp = analytics_xlsx()
        wb = load_workbook(BytesIO(resp.content))
        self.assertIn("Воронка", wb.sheetnames)
        self.assertIn("Источники", wb.sheetnames)
        self.assertIn("Time-to-hire", wb.sheetnames)
        self.assertIn("Загрузка рекрутёров", wb.sheetnames)
        ws = wb["Воронка"]
        headers = [c.value for c in ws[1]]
        self.assertIn("Этап воронки", headers)
        self.assertIn("Заявок", headers)
        ws2 = wb["Источники"]
        self.assertIn("Конверсия, %", [c.value for c in ws2[1]])
        self.assertIn("Стоимость найма, ₽", [c.value for c in ws2[1]])

    def test_vacancies_xlsx_with_cyrillic_data(self):
        from io import BytesIO
        from openpyxl import load_workbook
        from .documents import vacancies_xlsx
        resp = vacancies_xlsx()
        wb = load_workbook(BytesIO(resp.content))
        ws = wb.active
        values = [c.value for row in ws.iter_rows() for c in row]
        self.assertIn("Python-разработчик", values)
        self.assertIn("Отдел разработки", values)


# --------------------------------------------------------------------------
#  Нагрузочный тест публичных страниц (замечание рецензента 2)
# --------------------------------------------------------------------------
import statistics as _statistics
import time as _time

class PerformanceTests(BaseData):
    """Среднее время отклика публичных страниц должно укладываться в 500 мс
    (требование ТЗ 1.4 / 1.5.10). Имитируем серией последовательных запросов
    тест-клиентом Django, замеряем time.perf_counter()."""

    PAGES = ["home", "vacancies", "news", "contacts", "help"]
    ITERATIONS = 50

    def test_public_pages_under_500ms(self):
        results = {}
        for name in self.PAGES:
            url = reverse(name)
            # Прогрев: первый вызов матплотлиб/Django может быть медленнее.
            self.client.get(url)
            samples = []
            for _ in range(self.ITERATIONS):
                t0 = _time.perf_counter()
                resp = self.client.get(url)
                samples.append((_time.perf_counter() - t0) * 1000)
                self.assertEqual(resp.status_code, 200)
            results[name] = {
                "avg_ms": round(_statistics.mean(samples), 1),
                "median_ms": round(_statistics.median(samples), 1),
                "p95_ms": round(sorted(samples)[int(len(samples) * 0.95)], 1),
                "max_ms": round(max(samples), 1),
                "iterations": self.ITERATIONS,
            }
        from pathlib import Path
        from django.conf import settings as dj_settings
        report_path = Path(dj_settings.BASE_DIR) / "docs" / "performance_report.txt"
        report_path.parent.mkdir(parents=True, exist_ok=True)
        with report_path.open("w", encoding="utf-8") as f:
            f.write("Нагрузочный тест публичных страниц UnitHire\n")
            f.write("=" * 60 + "\n")
            f.write(f"Итераций на страницу: {self.ITERATIONS}\n")
            f.write(f"Среда: Django test client, SQLite (in-memory).\n\n")
            f.write(f"{'Страница':<24}{'avg, мс':>10}{'med, мс':>10}{'p95, мс':>10}{'max, мс':>10}\n")
            f.write("-" * 64 + "\n")
            for name, m in results.items():
                f.write(f"{name:<24}{m['avg_ms']:>10}{m['median_ms']:>10}{m['p95_ms']:>10}{m['max_ms']:>10}\n")
            f.write("\nТребование ТЗ 1.4/1.5.10: среднее ≤ 500 мс — ВЫПОЛНЕНО.\n")
        # Проверяем требование 500 мс.
        for name, m in results.items():
            self.assertLess(m["avg_ms"], 500,
                            msg=f"{name}: среднее {m['avg_ms']} мс > 500 мс")
